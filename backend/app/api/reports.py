import json
import tempfile
from pathlib import Path

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, status
from fastapi.responses import FileResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.templates import resolve_template_bytes, write_temp_docx
from app.core.deps import get_current_user
from app.database import get_db
from app.models import SurveyModule, SurveyRecord, User, UserRole
from app.schemas.survey import GenerateReportRequest


router = APIRouter(prefix="/api/reports", tags=["reports"])


def _remove_file(path: str) -> None:
    Path(path).unlink(missing_ok=True)


def _response_value(responses: dict, key: str) -> str:
    value = responses.get(key) or responses.get(key.title()) or ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


@router.post("/generate")
async def generate_report(
    body: GenerateReportRequest,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
) -> FileResponse:
    result = await db.execute(select(SurveyRecord).where(SurveyRecord.id.in_(body.record_ids)))
    records = list(result.scalars().all())
    if len(records) != len(set(body.record_ids)):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="One or more survey records were not found")
    if user.role == UserRole.surveyor and any(record.surveyor_id != user.id for record in records):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Cannot report on another surveyor's records")

    output = tempfile.NamedTemporaryFile(prefix="gdrpl-report-", suffix=".docx", delete=False)
    output.close()

    primary = records[0]
    category = (primary.structure_category or "").lower()
    module = (
        SurveyModule.utility_shifting
        if "utility" in category
        else SurveyModule.structure_inventory
    )

    docx_bytes, template_path = await resolve_template_bytes(db, module=module)
    temp_tpl: str | None = None
    if docx_bytes and (template_path is None or not template_path.is_file()):
        temp_tpl = write_temp_docx(docx_bytes)
        template_path = Path(temp_tpl)

    if template_path is not None and template_path.is_file():
        from docxtpl import DocxTemplate

        document = DocxTemplate(str(template_path))
        observations = "\n\n".join(
            f"{r.chainage}: {_response_value(r.responses_json, 'observations_recommendations') or _response_value(r.responses_json, 'observations') or '—'}"
            for r in records
        )
        recommendations = "\n\n".join(
            f"{r.chainage}: {_response_value(r.responses_json, 'recommendations') or _response_value(r.responses_json, 'observations_recommendations') or '—'}"
            for r in records
        )
        document.render(
            {
                "project_name": "",
                "chainage": ", ".join(r.chainage for r in records),
                "structure_category": primary.structure_category,
                "observations": observations,
                "recommendations": recommendations,
                "records": records,
                "record_count": len(records),
            }
        )
        document.save(output.name)
    else:
        from docx import Document

        document = Document()
        document.add_heading("GDRPL Survey Report", 0)
        for record in records:
            document.add_heading(f"{record.structure_category} — {record.chainage}", level=1)
            document.add_paragraph(f"Status: {record.status.value}")
            document.add_paragraph(
                f"Observations: {_response_value(record.responses_json, 'observations') or 'Not provided'}"
            )
            document.add_paragraph(
                f"Recommendations: {_response_value(record.responses_json, 'recommendations') or 'Not provided'}"
            )
        document.save(output.name)

    if temp_tpl:
        background_tasks.add_task(_remove_file, temp_tpl)
    background_tasks.add_task(_remove_file, output.name)
    return FileResponse(
        output.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="gdrpl-survey-report.docx",
        background=background_tasks,
    )
