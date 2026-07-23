"""Build editable GDRPL work-report DOCX: Q&A page + photo pages per structure."""

from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import UUID

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.models import Project, SurveyPhoto, SurveyRecord


SKIP_KEYS = {"gps", "capturedAt", "structure_category", "photos"}


def _fmt(value: Any) -> str:
    if value is None:
        return "—"
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def _category_label(key: str | None) -> str:
    if not key:
        return "Structure"
    return key.replace("_", " ").title()


def _response_rows(record: SurveyRecord) -> list[tuple[str, str]]:
    responses = record.responses_json or {}
    gps = responses.get("gps") if isinstance(responses.get("gps"), dict) else {}
    lat = gps.get("latitude", record.latitude)
    lon = gps.get("longitude", record.longitude)
    coords = f"{lat}, {lon}" if lat is not None and lon is not None else "—"

    rows: list[tuple[str, str]] = [
        ("Name of Road / Project", str(responses.get("name_of_road") or "")),
        (
            "Location of structure in Km.",
            f"{record.chainage or '—'} — {_category_label(record.structure_category).upper()}",
        ),
        ("Coordinates", coords),
        ("Structure Category", _category_label(record.structure_category)),
        ("Chainage", record.chainage or "—"),
        ("Status", record.status.value if hasattr(record.status, "value") else str(record.status)),
    ]

    for key, value in responses.items():
        if key in SKIP_KEYS:
            continue
        label = key.replace("_", " ").strip().title()
        rows.append((label, _fmt(value)))

    # Drop empty "Name of Road" if blank — keep row but show em dash
    fixed: list[tuple[str, str]] = []
    for label, data in rows:
        fixed.append((label, data if data not in ("", "None") else "—"))
    return fixed


def _set_cell_text(cell, text: str, *, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = bold
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_bordered_table(document: Document, headers: list[str], rows: list[list[str]], col_widths: list[Cm]):
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], header, bold=True, center=True)
        table.rows[0].cells[i].width = col_widths[i]
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            _set_cell_text(table.rows[r_idx].cells[c_idx], value, center=(c_idx == 0))
            table.rows[r_idx].cells[c_idx].width = col_widths[c_idx]
    return table


def _photo_paths(photos: list[SurveyPhoto]) -> list[Path]:
    paths: list[Path] = []
    for photo in photos:
        p = Path(photo.local_path) if photo.local_path else None
        if p and p.is_file():
            paths.append(p)
    return paths


def _add_photo_pages(document: Document, project_name: str, structure_index: int, photos: list[Path]) -> None:
    """2x2 photo grids; automatically adds more pages when photo count > 4."""
    if not photos:
        document.add_page_break()
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(project_name or "GDRPL Survey")
        run.bold = True
        run.font.size = Pt(14)
        document.add_paragraph(f"Page-2 (Photos of Structure-{structure_index})")
        document.add_paragraph("No photos captured for this structure.")
        return

    chunk_size = 4
    page_no = 2
    for start in range(0, len(photos), chunk_size):
        chunk = photos[start : start + chunk_size]
        document.add_page_break()
        header = document.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(project_name or "GDRPL Survey")
        run.bold = True
        run.font.size = Pt(14)

        label = document.add_paragraph()
        label_run = label.add_run(f"Page-{page_no} (Photos of Structure-{structure_index})")
        label_run.bold = True

        # Always a 2x2 grid; empty cells if fewer than 4 on last page
        table = document.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        for i in range(4):
            cell = table.rows[i // 2].cells[i % 2]
            cell.width = Cm(8)
            if i < len(chunk):
                try:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cell.paragraphs[0].add_run()
                    run.add_picture(str(chunk[i]), width=Inches(2.8))
                except Exception:
                    _set_cell_text(cell, f"Photo-{start + i + 1} (unavailable)", center=True)
            else:
                _set_cell_text(cell, f"Photo-{start + i + 1}", center=True)
        page_no += 1


def build_work_report_docx(
    *,
    project_name: str,
    records: list[SurveyRecord],
) -> bytes:
    """Return editable .docx bytes for the work report layout."""
    document = Document()

    # Narrow margins for report look
    for section in document.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    for index, record in enumerate(records, start=1):
        if index > 1:
            document.add_page_break()

        header = document.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(project_name or "GDRPL Survey")
        run.bold = True
        run.font.size = Pt(14)

        page_title = document.add_paragraph()
        page_title_run = page_title.add_run(f"Page-1 (Structure-{index})")
        page_title_run.bold = True

        data_rows = _response_rows(record)
        table_rows = [[str(i), desc, data] for i, (desc, data) in enumerate(data_rows, start=1)]
        _add_bordered_table(
            document,
            ["Sr. No", "Description", "Data"],
            table_rows,
            [Cm(2), Cm(7.5), Cm(7.5)],
        )

        photos = _photo_paths(list(getattr(record, "photos", []) or []))
        _add_photo_pages(document, project_name, index, photos)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def load_records_for_report(db: AsyncSession, record_ids: list[UUID]) -> tuple[str, list[SurveyRecord]]:
    result = await db.execute(
        select(SurveyRecord)
        .where(SurveyRecord.id.in_(record_ids))
        .options(selectinload(SurveyRecord.photos), selectinload(SurveyRecord.project))
    )
    records = list(result.scalars().unique().all())
    # Preserve request order
    by_id = {r.id: r for r in records}
    ordered = [by_id[i] for i in record_ids if i in by_id]
    project_name = ""
    if ordered:
        project = ordered[0].project
        if project is None and ordered[0].project_id:
            project = await db.get(Project, ordered[0].project_id)
        project_name = project.name if project else ""
    return project_name, ordered
